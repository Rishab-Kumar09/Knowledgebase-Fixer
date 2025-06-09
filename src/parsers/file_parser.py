import logging
from pathlib import Path
from typing import Dict, List, Union, Optional
import frontmatter
from bs4 import BeautifulSoup
import markdown
from PyPDF2 import PdfReader
from docx import Document
import os
import json
from datetime import datetime
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

class DateTimeEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, datetime):
            return obj.isoformat()
        return super().default(obj)

@dataclass
class Article:
    path: str
    content: str
    metadata: Dict = field(default_factory=dict)
    file_type: str = ''

class FileParser:
    """Parser for handling different types of KB article files."""
    
    def __init__(self):
        self.supported_extensions = {'.md', '.html', '.txt', '.pdf', '.doc', '.docx'}
    
    def parse_directory(self, directory: Union[str, Path]) -> List[Article]:
        """Parse all supported files in the given directory."""
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        articles = []
        for file_path in directory.rglob('*'):
            if file_path.suffix.lower() in self.supported_extensions:
                try:
                    article = self._parse_file(file_path)
                    if article:
                        articles.append(article)
                except Exception as e:
                    logger.error(f"Error parsing {file_path}: {str(e)}")
        
        return articles
    
    def _parse_file(self, file_path: Path) -> Optional[Article]:
        """Parse a single file based on its extension."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension in ['.doc', '.docx']:
            return self._parse_word(file_path)
        elif extension == '.md':
            return self._parse_markdown(file_path)
        elif extension == '.html':
            return self._parse_html(file_path)
        else:  # .txt
            return self._parse_text(file_path)
    
    def _parse_pdf(self, file_path: Path) -> Optional[Article]:
        """Parse PDF content."""
        try:
            reader = PdfReader(file_path)
            text = ""
            metadata = {}
            
            # Extract metadata
            if reader.metadata:
                for key, value in reader.metadata.items():
                    if value:
                        # Remove the leading '/' from PDF metadata keys
                        clean_key = key[1:] if key.startswith('/') else key
                        metadata[clean_key] = str(value)
            
            # Extract text from all pages
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return Article(
                path=str(file_path),
                content=text.strip(),
                metadata=metadata,
                file_type='pdf'
            )
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            return None
    
    def _parse_word(self, file_path: Path) -> Optional[Article]:
        """Parse Word document content."""
        try:
            doc = Document(file_path)
            text = ""
            metadata = {
                'title': os.path.splitext(file_path.name)[0],
                'author': '',
                'created': '',
                'modified': ''
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Try to extract core properties
            try:
                core_props = doc.core_properties
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.created:
                    metadata['created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['modified'] = core_props.modified.isoformat()
            except:
                pass  # Skip if core properties are not accessible
            
            return Article(
                path=str(file_path),
                content=text.strip(),
                metadata=metadata,
                file_type='word'
            )
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            return None
    
    def _parse_markdown(self, file_path: Path) -> Optional[Article]:
        """Parse markdown content with frontmatter."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse frontmatter if present
            post = frontmatter.loads(content)
            metadata = dict(post.metadata)
            
            # Convert markdown to HTML for consistent processing
            html_content = markdown.markdown(post.content)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata from first h1 tag if present
            title = soup.find('h1')
            metadata['title'] = title.text if title else os.path.basename(file_path)
            metadata['created_at'] = datetime.fromtimestamp(os.path.getctime(file_path))
            metadata['updated_at'] = datetime.fromtimestamp(os.path.getmtime(file_path))
            
            return Article(
                path=str(file_path),
                content=content,
                metadata=metadata,
                file_type='markdown'
            )
        except Exception as e:
            logger.error(f"Error parsing markdown {file_path}: {str(e)}")
            return None
    
    def _parse_html(self, file_path: Path) -> Optional[Article]:
        """Parse HTML content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract metadata from title tag and meta tags
            title_tag = soup.find('title')
            metadata = {
                'title': title_tag.text if title_tag else os.path.basename(file_path),
                'created_at': datetime.fromtimestamp(os.path.getctime(file_path)),
                'updated_at': datetime.fromtimestamp(os.path.getmtime(file_path))
            }
            
            # Extract meta tags
            for meta in soup.find_all('meta'):
                name = meta.get('name', '').lower()
                content = meta.get('content', '')
                if name and content:
                    metadata[name] = content
            
            return Article(
                path=str(file_path),
                content=str(soup.body) if soup.body else content,
                metadata=metadata,
                file_type='html'
            )
        except Exception as e:
            logger.error(f"Error parsing HTML {file_path}: {str(e)}")
            return None
    
    def _parse_text(self, file_path: Path) -> Optional[Article]:
        """Parse plain text content."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # For text files, use filename as title and file system metadata
        metadata = {
            'title': os.path.basename(file_path),
            'created_at': datetime.fromtimestamp(os.path.getctime(file_path)),
            'updated_at': datetime.fromtimestamp(os.path.getmtime(file_path))
        }
        
        return Article(
            path=str(file_path),
            content=content,
            metadata=metadata,
            file_type='text'
        )

    def create_article_from_dict(self, data: Dict) -> Article:
        """
        Create an Article object from a dictionary containing article data.
        Useful for creating articles from database records.
        
        Args:
            data (Dict): Dictionary containing article data with keys:
                - path: str
                - content: str
                - metadata: Dict (optional)
                - file_type: str (optional)
        
        Returns:
            Article: An Article object created from the dictionary data
        """
        return Article(
            path=data['path'],
            content=data['content'],
            metadata=data.get('metadata', {}),
            file_type=data.get('file_type', 'text')
        ) 